import os
import io
import re
import json
import datetime
import collections
import streamlit as st
import pdfplumber
import requests
from docx import Document

# ---------------- Config ----------------
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL_NAME = "gpt-4o-mini"

PLACEHOLDER_RE = re.compile(r"\[([A-Za-z0-9_]+)\]")

# ---------------- Helpers ----------------


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    try:
        parts = []
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n\n".join(parts)
    except Exception:
        return ""


def find_placeholders(doc: Document) -> list:
    found = set()

    def scan(paragraphs):
        for p in paragraphs:
            for k in PLACEHOLDER_RE.findall(p.text):
                found.add(k)

    scan(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                scan(cell.paragraphs)

    try:
        for s in doc.sections:
            scan(s.header.paragraphs)
            scan(s.footer.paragraphs)
    except Exception:
        pass

    return sorted(found)


def extract_first_json_block(text: str) -> str:
    start = text.find("{")
    if start == -1:
        return ""
    braces = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            braces += 1
        elif text[i] == "}":
            braces -= 1
            if braces == 0:
                return text[start:i + 1]
    return ""


def call_openrouter(messages: list, max_tokens: int = 1200) -> str:
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OPENROUTER_API_KEY environment variable not set.")

    payload = {
        "model": MODEL_NAME,
        "messages": messages,
        "temperature": 0.0,
        "max_tokens": max_tokens
    }
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}",
               "Content-Type": "application/json"}
    r = requests.post(OPENROUTER_URL, headers=headers,
                      json=payload, timeout=120)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

# ---------------- Extraction Modes ----------------


def llm_high_accuracy(placeholders: list, pdf_text: str) -> dict:
    system = (
        "You are an expert insurance GLR extractor. Produce accurate concise values for each placeholder. "
        "Normalize dates (MM/DD/YYYY). Prefer explicit evidence. Don't hallucinate contradictions."
    )

    user = (
        f"Placeholders:\n{json.dumps(placeholders, indent=2)}\n\n"
        f"PDF TEXT:\n----BEGIN----\n{pdf_text}\n----END----\n\n"
        "Return ONLY JSON:\n{ \"mapping\": {...}, \"confidences\": {...} }"
    )

    out = call_openrouter(
        [{"role": "system", "content": system},
            {"role": "user", "content": user}],
        max_tokens=1500
    )

    block = extract_first_json_block(out)
    parsed = json.loads(block if block else out)

    mapping = parsed.get("mapping", {})
    confidences = parsed.get("confidences", {})

    for ph in placeholders:
        mapping.setdefault(ph, "")
        confidences.setdefault(ph, 0.0)

    return {"mapping": mapping, "confidences": confidences}


def llm_strict_validation(placeholders: list, pdf_text: str) -> dict:
    system = (
        "You are a strict validator. Mark each field as VERIFIED, INFERRED, or MISSING. "
        "Provide evidence for each field."
    )

    user = (
        f"Placeholders:\n{json.dumps(placeholders, indent=2)}\n\n"
        f"PDF TEXT:\n----BEGIN----\n{pdf_text}\n----END----\n\n"
        "Return ONLY JSON: { mapping:{}, status:{}, reasons:{} }"
    )

    out = call_openrouter(
        [{"role": "system", "content": system},
            {"role": "user", "content": user}],
        max_tokens=1500
    )

    block = extract_first_json_block(out)
    parsed = json.loads(block if block else out)

    for ph in placeholders:
        parsed.setdefault("mapping", {}).setdefault(ph, "")
        parsed.setdefault("status", {}).setdefault(ph, "MISSING")
        parsed.setdefault("reasons", {}).setdefault(ph, "")

    return parsed


def llm_field_audit(placeholders: list, pdf_text: str) -> dict:
    system = "Extract each placeholder value and provide 1–2 sentence evidence."

    user = (
        f"Placeholders:\n{json.dumps(placeholders, indent=2)}\n\n"
        f"PDF TEXT:\n----BEGIN----\n{pdf_text}\n----END----\n\n"
        "Return ONLY JSON: {\"field_values\":{}, \"evidence\":{}}"
    )

    out = call_openrouter(
        [{"role": "system", "content": system},
            {"role": "user", "content": user}],
        max_tokens=1500
    )

    block = extract_first_json_block(out)
    parsed = json.loads(block if block else out)

    for ph in placeholders:
        parsed.setdefault("field_values", {}).setdefault(ph, "")
        parsed.setdefault("evidence", {}).setdefault(ph, "")

    return parsed


def llm_voting(placeholders: list, pdf_text: str, rounds: int = 3) -> dict:
    candidates = {ph: [] for ph in placeholders}
    temps = [0.0, 0.2, 0.6][:rounds]

    for t in temps:
        system = f"You are an extractor (voting pass temp={t}). Return ONLY JSON mapping placeholders->values."
        user = (
            f"Placeholders:\n{json.dumps(placeholders, indent=2)}\n\n"
            f"PDF TEXT:\n----BEGIN----\n{pdf_text}\n----END----\n"
            "Return ONLY raw JSON."
        )

        payload = {
            "model": MODEL_NAME,
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": user}],
            "temperature": t,
            "max_tokens": 1200
        }
        headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}",
                   "Content-Type": "application/json"}
        r = requests.post(OPENROUTER_URL, headers=headers,
                          json=payload, timeout=120)
        r.raise_for_status()

        out = r.json()["choices"][0]["message"]["content"]
        block = extract_first_json_block(out)

        try:
            parsed = json.loads(block if block else out)
        except Exception:
            continue

        for ph in placeholders:
            if parsed.get(ph):
                candidates[ph].append(parsed[ph].strip())

    final_map = {}
    votes = {}

    for ph, vals in candidates.items():
        if not vals:
            final_map[ph] = ""
            votes[ph] = []
        else:
            counter = collections.Counter(vals)
            winner = counter.most_common(1)[0][0]
            final_map[ph] = winner
            votes[ph] = dict(counter)

    return {"mapping": final_map, "votes": votes, "candidates": candidates}

# ---------------- Fallback Heuristics ----------------


def fallback_structured(pdf_text: str, placeholders: list) -> dict:
    data = {k: "" for k in placeholders}

    m = re.search(
        r"(Insured|Insured Name|Policyholder)[^\n]{0,20}([A-Z][A-Za-z .,'-]{2,80})", pdf_text, re.I)
    if m:
        data["INSURED_NAME"] = m.group(2).strip()

    addr = re.search(
        r"([0-9]{1,5}[^,\n]+),\s*([A-Za-z\s]+),\s*([A-Z]{2})\s*(\d{5})", pdf_text)
    if addr:
        data["INSURED_H_STREET"] = addr.group(1).strip()
        data["INSURED_H_CITY"] = addr.group(2).strip()
        data["INSURED_H_STATE"] = addr.group(3).strip()
        data["INSURED_H_ZIP"] = addr.group(4).strip()

    md = re.search(r"([0-1]?\d[\/\-][0-3]?\d[\/\-]\d{2,4})", pdf_text)
    if md:
        data["DATE_INSPECTED"] = md.group(1)

    m2 = re.search(r"(Mortgagee|Lender|Bank)[^\n]{0,20}(.+)", pdf_text, re.I)
    if m2:
        data["MORTGAGEE"] = m2.group(2).strip()

    if "DATE_RECEIVED" in data:
        data["DATE_RECEIVED"] = datetime.datetime.now().strftime("%m/%d/%Y")

    return data

# ---------------- DOCX Fill ----------------


def fill_docx(template_bytes: bytes, mapping: dict) -> bytes:
    doc = Document(io.BytesIO(template_bytes))

    def repl(p):
        txt = p.text
        for k, v in mapping.items():
            token = f"[{k}]"
            if token in txt:
                txt = txt.replace(token, str(v))
        p.text = txt

    for p in doc.paragraphs:
        repl(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl(p)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------- Streamlit UI ----------------


st.set_page_config(page_title="GLR Advanced Modes", layout="wide")
st.title("GLR — Advanced Extraction Modes (High-accuracy, Strict, Audit, Voting)")

uploaded_template = st.file_uploader("Upload Template (.docx)", type=["docx"])
uploaded_pdfs = st.file_uploader(
    "Upload Photo Report PDFs (text-based)", type=["pdf"], accept_multiple_files=True)

mode = st.selectbox("Select extraction mode", [
                    "high_accuracy", "strict_validation", "field_audit", "multi_llm_voting"])

if not uploaded_template:
    st.info("Please upload a template (.docx)")
    st.stop()

if not uploaded_pdfs:
    st.info("Please upload at least one PDF")
    st.stop()

template_bytes = uploaded_template.read()

# Find placeholders
doc_preview = Document(io.BytesIO(template_bytes))
placeholders = find_placeholders(doc_preview)
st.subheader("Detected placeholders")
st.write(placeholders)

# Extract PDFs
combined = []
for f in uploaded_pdfs:
    txt = extract_text_from_pdf_bytes(f.read())
    if not txt.strip():
        st.warning(f"No text found in {f.name} — skipped (not OCR version).")
    else:
        combined.append(txt)

combined_text = "\n\n".join(combined)
st.text_area("Extracted Text", combined_text[:20000], height=300)

# Extraction
if st.button("Run Extraction"):
    with st.spinner("Extracting…"):
        fallback = fallback_structured(combined_text, placeholders)

        if mode == "high_accuracy":
            result = llm_high_accuracy(placeholders, combined_text)
            mapping = result["mapping"]
            extras = result
        elif mode == "strict_validation":
            result = llm_strict_validation(placeholders, combined_text)
            mapping = result["mapping"]
            extras = result
        elif mode == "field_audit":
            result = llm_field_audit(placeholders, combined_text)
            mapping = result["field_values"]
            extras = result
        else:
            result = llm_voting(placeholders, combined_text)
            mapping = result["mapping"]
            extras = result

        # Fill blanks with fallback
        for ph in placeholders:
            if not mapping.get(ph):
                mapping[ph] = fallback.get(ph, "")

        # Try no-missing-data inference
        still_empty = [k for k, v in mapping.items() if v == ""]
        if still_empty:
            system_nm = "Infer missing insurance fields. NEVER output empty or N/A."
            user_nm = (
                f"Fill these fields: {still_empty}\n"
                f"PDF TEXT:\n{combined_text}\n"
                "Return ONLY JSON."
            )
            try:
                nm_out = call_openrouter(
                    [{"role": "system", "content": system_nm},
                        {"role": "user", "content": user_nm}],
                    max_tokens=800
                )
                block = extract_first_json_block(nm_out)
                nm_map = json.loads(block if block else nm_out)
                for k, v in nm_map.items():
                    if mapping[k] == "":
                        mapping[k] = v
            except:
                pass

        st.success("Extraction complete!")
        st.json(mapping)

        out_bytes = fill_docx(template_bytes, mapping)
        st.download_button(
            "Download Completed .docx",
            data=out_bytes,
            file_name=f"filled_glr_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
